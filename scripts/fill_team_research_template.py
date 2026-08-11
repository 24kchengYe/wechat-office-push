#!/usr/bin/env python
"""Fill the BCL long-form team-research template from a structured DOCX.

This automation only saves a draft when ``--save-draft`` is supplied. It
never invokes the publish button.
"""

from __future__ import annotations

import argparse
import base64
import copy
import gzip
import html
import json
import re
import time
from datetime import datetime
from pathlib import Path

from bs4 import BeautifulSoup, Tag
from docx import Document

import fill_wechat_template as ui


def clean(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").split())


def after_label(value: str) -> str:
    return value.split("】", 1)[1].strip() if "】" in value else value.strip()


def parse_docx(path: Path, images_dir: Path) -> dict:
    document = Document(path)
    paragraphs = [clean(item.text) for item in document.paragraphs]
    styles = [item.style.name for item in document.paragraphs]

    metadata = {}
    for text in paragraphs[:12]:
        if text.startswith("【论文题目】"):
            metadata["title_en"] = after_label(text)
        elif text.startswith("【作者】"):
            metadata["authors"] = after_label(text)
        elif text.startswith(("a.", "b.", "c.")):
            metadata.setdefault("affiliations", []).append(text)
        elif text.startswith("【通讯作者邮箱】"):
            metadata["email"] = after_label(text)
        elif text.startswith("【期刊信息】"):
            metadata["journal"] = after_label(text)
        elif text.startswith("【全文链接】"):
            metadata["doi_url"] = after_label(text)

    headings = []
    for index, (text, style) in enumerate(zip(paragraphs, styles)):
        if text and style == "Heading 1":
            headings.append((index, text))

    section_ranges = {}
    for position, (start, title) in enumerate(headings):
        end = headings[position + 1][0] if position + 1 < len(headings) else len(paragraphs)
        compact_title = re.sub(r"\s+", "", title)
        if "摘要" in compact_title:
            key = "guide"
        elif "方法" in compact_title:
            key = "methods"
        elif "结果" in compact_title:
            key = "results"
        elif "意义" in compact_title:
            key = "significance"
        else:
            continue
        section_ranges[key] = (start + 1, end)

    figures = sorted(images_dir.glob("fig*.png"), key=lambda p: int(re.search(r"\d+", p.stem).group()))
    if len(figures) != 5:
        raise ValueError(f"expected five figures, got {len(figures)}")

    def items(key: str) -> list[dict]:
        start, end = section_ranges[key]
        output = []
        figure_index = 0 if key == "methods" else 2 if key == "results" else None
        for index in range(start, end):
            text = paragraphs[index]
            if not text or text.startswith("关键词："):
                continue
            style = styles[index]
            if style == "Heading 2":
                output.append({"type": "subheading", "text": text})
            elif style == "Caption":
                if figure_index is None or figure_index >= len(figures):
                    raise ValueError(f"unexpected caption in {key}: {text}")
                output.append(
                    {
                        "type": "figure",
                        "path": figures[figure_index],
                        "caption": text,
                    }
                )
                figure_index += 1
            else:
                output.append({"type": "paragraph", "text": text})
        return output

    doi_url = metadata["doi_url"]
    citation = ""
    if document.tables:
        citation = clean(document.tables[0].cell(0, 0).text).removeprefix("原文信息：")
    citation_full = citation
    if doi_url and doi_url not in citation_full:
        citation_full = (citation_full + " " + doi_url).strip()
    return {
        "headline": paragraphs[0],
        **metadata,
        "doi": doi_url.removeprefix("https://doi.org/"),
        "citation": citation,
        "citation_full": citation_full,
        "guide": items("guide"),
        "methods": items("methods"),
        "results": items("results"),
        "significance": items("significance"),
    }


def paragraph(text: str) -> str:
    return (
        '<p style="margin:0 0 16px;text-align:justify;line-height:1.8em;">'
        '<span style="font-size:15px;color:#3f3f3f;letter-spacing:0.5px;">'
        + html.escape(text)
        + "</span></p>"
    )


def subheading(text: str) -> str:
    return (
        '<p style="margin:22px 0 12px;line-height:1.8em;">'
        '<strong style="font-size:16px;color:#5c3a62;">'
        + html.escape(text)
        + "</strong></p>"
    )


def figure(path: Path, caption: str) -> str:
    data = base64.b64encode(path.read_bytes()).decode("ascii")
    return (
        '<p style="margin:18px 0 8px;text-align:center;">'
        f'<img class="rich_pages wxw-img js_insertlocalimg" data-type="png" '
        f'src="data:image/png;base64,{data}" '
        'style="display:block;width:100%;height:auto!important;"/>'
        "</p>"
        '<p style="margin:0 0 18px;text-align:center;line-height:1.6em;">'
        '<span style="font-size:13px;color:#777777;">'
        + html.escape(caption)
        + "</span></p>"
    )


def content_block(items: list[dict]) -> str:
    chunks = ['<section style="margin:0 12px;">']
    for item in items:
        if item["type"] == "paragraph":
            chunks.append(paragraph(item["text"]))
        elif item["type"] == "subheading":
            chunks.append(subheading(item["text"]))
        elif item["type"] == "figure":
            chunks.append(figure(item["path"], item["caption"]))
    chunks.append("</section>")
    return "".join(chunks)


def display_authors(value: str) -> str:
    author_names = re.sub(r"([A-Za-z]+)[abc](?=,|$)", r"\1", value)
    return author_names.replace("*c", "*")


def metadata_html(article: dict) -> str:
    author_names = display_authors(article["authors"])
    rows = [
        ("【作者】", author_names),
        *[("", affiliation) for affiliation in article["affiliations"]],
        ("【通讯作者邮箱】", article["email"]),
        ("【期刊信息】", article["journal"]),
        ("【原文信息】", article["citation_full"]),
    ]
    output = []
    for label, value in rows:
        output.append(
            '<p style="margin:0 0 8px;line-height:1.7em;text-align:left;">'
            + (f"<strong>{html.escape(label)}</strong>" if label else "")
            + f'<span style="font-size:14px;color:#3f3f3f;">{html.escape(value)}</span>'
            + "</p>"
        )
    return "".join(output)


def template_fragment(asset: Path, article: dict) -> str:
    raw = gzip.decompress(base64.b64decode(asset.read_text(encoding="ascii"))).decode(
        "utf-8", "replace"
    )
    match = re.search(r"<!--StartFragment-->(.*)<!--EndFragment-->", raw, re.S)
    soup = BeautifulSoup(match.group(1) if match else raw, "html.parser")
    roots = [node for node in soup.contents if isinstance(node, Tag)]
    if len(roots) < 23:
        raise ValueError("team-research template structure changed")

    heading_guide = str(roots[6])
    heading_methods = str(roots[8]).replace("研究背景", "研究方法")
    heading_results = str(roots[10]).replace("研究结果", "研究结果与讨论")
    heading_significance = str(roots[12]).replace("基金资助", "研究意义")
    year_match = re.search(r"\b(20\d{2})\b", article.get("journal", ""))
    publication_year = year_match.group(1) if year_match else str(datetime.now().year)
    year_art = re.sub(r"20\d{2}", publication_year, str(roots[15]), count=1)

    return "".join(
        [
            metadata_html(article),
            str(roots[5]),
            heading_guide,
            content_block(article["guide"]),
            heading_methods,
            content_block(article["methods"]),
            heading_results,
            content_block(article["results"]),
            heading_significance,
            content_block(article["significance"]),
            str(roots[14]),
            year_art,
            str(roots[16]),
            str(roots[17]),
            str(roots[19]),
            str(roots[20]),
            str(roots[21]),
            str(roots[22]),
        ]
    )


def recommendation(article: dict) -> str:
    guide_text = next(
        (item["text"] for item in article.get("guide", []) if item["type"] == "paragraph"),
        "",
    )
    value = clean("。".join(part.strip("。") for part in [article["headline"], guide_text] if part))
    if ui.recommendation_count(value) <= 120:
        return value
    shortened = ""
    for character in value:
        candidate = shortened + character
        if ui.recommendation_count(candidate + "…") > 120:
            break
        shortened = candidate
    return shortened.rstrip("，。；： ") + "…"


def rich_editors(automation, root):
    all_elements = root.FindAll(ui.TREE_DESCENDANTS, automation.CreateTrueCondition())
    body = title = None
    for index in range(all_elements.Length):
        element = all_elements.GetElement(index)
        bounds = element.CurrentBoundingRectangle
        class_name = element.CurrentClassName
        if bounds.left < -1300 or "ProseMirror" not in class_name:
            continue
        height = bounds.bottom - bounds.top
        if height > 1000 and bounds.right - bounds.left > 500:
            body = element
        elif height < 180 and bounds.right - bounds.left > 400:
            title = element
    if not body or not title:
        raise RuntimeError("BCL team-research title/body editors were not found")
    return title, body


def select_codex_task(UIA, automation, root, hwnd: int, task_name: str):
    matches = ui.find_all(automation, root, ui.UIA_NAME, task_name)
    for index in range(matches.Length):
        candidate = matches.GetElement(index)
        if candidate.CurrentControlType == ui.UIA_BUTTON:
            try:
                ui.invoke(UIA, candidate)
                time.sleep(0.6)
                return automation.ElementFromHandle(hwnd)
            except Exception:
                continue
    return root


def set_platform_fields(UIA, automation, root, hwnd: int, article: dict) -> tuple[str, str]:
    rec = recommendation(article)
    if ui.recommendation_count(rec) > 120:
        raise ValueError("platform recommendation exceeds 120 characters")
    description = ui.find_first(automation, root, ui.UIA_AUTOMATION_ID, "js_description")
    ui.scroll_into_view(UIA, description)
    ui.click_element(description, hwnd)
    ui.set_unicode_clipboard(rec)
    ui.chord("a")
    ui.chord("v")

    doi_url = article["doi_url"]
    url_area = ui.find_first(automation, root, ui.UIA_AUTOMATION_ID, "js_article_url_area")
    ui.scroll_into_view(UIA, url_area)
    texts = ui.find_all(automation, url_area, ui.UIA_CONTROL_TYPE, 50020)
    current_url = None
    for index in range(texts.Length):
        candidate = texts.GetElement(index)
        if candidate.CurrentName.startswith(("https://", "http://")):
            current_url = candidate
            break
    if not current_url:
        raise RuntimeError("current original link not found")
    if current_url.CurrentName != doi_url:
        ui.click_element(current_url, hwnd)
        url_input = ui.find_first(automation, root, ui.UIA_NAME, "输入或粘贴原文链接")
        url_input.GetCurrentPattern(ui.VALUE_PATTERN).QueryInterface(
            UIA.IUIAutomationValuePattern
        ).SetValue(doi_url)
        ui.invoke(UIA, ui.find_first(automation, root, ui.UIA_NAME, "确定"))
        time.sleep(1.5)
        root = automation.ElementFromHandle(hwnd)
        url_area = ui.find_first(automation, root, ui.UIA_AUTOMATION_ID, "js_article_url_area")
        saved_urls = ui.find_all(automation, url_area, ui.UIA_CONTROL_TYPE, 50020)
        if not any(saved_urls.GetElement(index).CurrentName == doi_url for index in range(saved_urls.Length)):
            raise RuntimeError("original link was not updated after confirmation")

    collection = "论文推荐"
    area = ui.find_first(automation, root, ui.UIA_AUTOMATION_ID, "js_article_tags_area")
    selected_global = ui.find_first(automation, root, ui.UIA_NAME, collection)
    # Older BCL editors expose only the visible label/value, not the newer
    # js_article_tags_area automation id. A visible exact value is sufficient
    # to validate the inherited collection from the published template.
    if not area and selected_global:
        return rec, doi_url
    if not area:
        raise RuntimeError("collection setting not found")
    if not ui.find_first(automation, area, ui.UIA_NAME, collection):
        unassigned = ui.find_first(automation, area, ui.UIA_NAME, "未添加")
        trigger = automation.RawViewWalker.GetParentElement(unassigned)
        trigger.GetCurrentPattern(10018).QueryInterface(
            UIA.IUIAutomationLegacyIAccessiblePattern
        ).DoDefaultAction()
        time.sleep(0.7)
        root = automation.ElementFromHandle(hwnd)
        search = ui.find_first(automation, root, ui.UIA_NAME, "请选择合集")
        ui.click_element(search, hwnd)
        ui.set_unicode_clipboard(collection)
        ui.chord("a")
        ui.chord("v")
        time.sleep(0.7)
        matches = ui.find_all(automation, root, ui.UIA_NAME, collection)
        option = next(
            (
                matches.GetElement(index)
                for index in range(matches.Length)
                if matches.GetElement(index).CurrentControlType == ui.UIA_LIST_ITEM
            ),
            None,
        )
        if not option:
            raise RuntimeError("paper-recommendation collection option not found")
        ui.invoke(UIA, option)
        ui.invoke(UIA, ui.find_first(automation, root, ui.UIA_NAME, "确认"))
        time.sleep(0.8)
    return rec, doi_url


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--images-dir", type=Path, required=True)
    parser.add_argument("--hwnd", type=int, required=True)
    parser.add_argument("--task-name", default="BCL公众号")
    parser.add_argument("--save-draft", action="store_true")
    parser.add_argument(
        "--settings-only",
        action="store_true",
        help="Keep the already validated title/body and finish metadata/save only.",
    )
    args = parser.parse_args()

    root_dir = Path(__file__).resolve().parent.parent
    article = parse_docx(args.docx.resolve(), args.images_dir.resolve())
    fragment = template_fragment(
        root_dir / "templates" / "bcl_team_research_template.cfhtml.gz.b64", article
    )

    UIA, automation, root = ui.uia_client(args.hwnd)
    root = select_codex_task(UIA, automation, root, args.hwnd, args.task_name)
    if args.settings_only:
        checks = {"title_body_previously_validated": True}
    else:
        title, body = rich_editors(automation, root)
        ui.scroll_into_view(UIA, title)
        ui.click_element(title, args.hwnd)
        ui.set_unicode_clipboard(article["headline"])
        ui.chord("a")
        ui.chord("v")
        ui.scroll_into_view(UIA, body)
        ui.click_element(body, args.hwnd)
        ui.set_rich_clipboard(fragment, article["guide"][0]["text"])
        ui.chord("a")
        ui.chord("v")
        time.sleep(15)

        title_text = ui.text_of(UIA, title).strip()
        body_text = ui.text_of(UIA, body)
        images = ui.find_all(automation, body, ui.UIA_CONTROL_TYPE, ui.UIA_IMAGE)
        paper_images = 0
        total_images = 0
        for index in range(images.Length):
            class_name = images.GetElement(index).CurrentClassName
            if "wxw-img" in class_name:
                total_images += 1
            if "js_insertlocalimg" in class_name:
                paper_images += 1
        checks = {
            "headline": title_text == article["headline"],
            "paper_title": article["title_en"] in body_text,
            "authors": display_authors(article["authors"]).split(",", 1)[0] in body_text,
            "doi": article["doi"] in body_text,
            "guide": article["guide"][0]["text"][:30] in body_text,
            "method": article["methods"][0]["text"] in body_text,
            "results": article["results"][0]["text"] in body_text,
            "significance": article["significance"][0]["text"][:30] in body_text,
            "old_title_removed": "中国城市空地识别及其分布规律研究" not in body_text,
            "images": total_images,
            "paper_images": paper_images,
        }
        boolean_keys = [key for key in checks if key not in {"images", "paper_images"}]
        if not all(checks[key] for key in boolean_keys) or total_images != 7 or paper_images != 5:
            raise RuntimeError("pre-save validation failed: " + json.dumps(checks, ensure_ascii=False))

    root = automation.ElementFromHandle(args.hwnd)
    rec, doi_url = set_platform_fields(UIA, automation, root, args.hwnd, article)
    saved = False
    if args.save_draft:
        root = automation.ElementFromHandle(args.hwnd)
        save = ui.find_first(automation, root, ui.UIA_NAME, "保存为草稿")
        ui.invoke(UIA, save)
        time.sleep(6)
        saved = True

    print(
        json.dumps(
            {
                "headline": article["headline"],
                "doi": article["doi"],
                "checks": checks,
                "platform_recommendation": rec,
                "platform_recommendation_count": ui.recommendation_count(rec),
                "original_link": doi_url,
                "collection": "论文推荐",
                "saved_as_draft": saved,
                "published": False,
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
